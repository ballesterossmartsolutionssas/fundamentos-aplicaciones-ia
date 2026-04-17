from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
REPORTS_DIR = ROOT / "reports"
INPUT_PATH = REPORTS_DIR / "informe_miniproyecto_2.md"
OUTPUT_PATH = REPORTS_DIR / "informe_miniproyecto_2.docx"
COURSE_NAME = "Fundamentos y Aplicaciones de Inteligencia Artificial"
UNIVERSITY_NAME = "Universidad Autonoma de Occidente"
PROFESSOR_NAME = "Juan Sebastian Mosquera Maturana"
PROJECT_NAME = "MiniProyecto 2 - Clasificacion de niveles de obesidad"
AUTHORS = [
    "Valentina Popo Montilla",
    "Juan Camilo Balleresteros Sierra",
    "Santigo Rodriguez Gacha",
]


def set_default_font(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal_style.font.size = Pt(11)


def add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def add_cover_page(document: Document) -> None:
    top = document.add_paragraph()
    top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = top.add_run(UNIVERSITY_NAME)
    run.bold = True
    run.font.size = Pt(16)

    course = document.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = course.add_run(COURSE_NAME)
    run.bold = True
    run.font.size = Pt(14)

    document.add_paragraph("")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(PROJECT_NAME)
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Informe final")
    run.italic = True
    run.font.size = Pt(13)

    document.add_paragraph("")
    document.add_paragraph("")

    info = [
        ("Profesor", PROFESSOR_NAME),
        ("Fecha", datetime.now().strftime("%d/%m/%Y")),
    ]
    for label, value in info:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"{label}: {value}")
        run.font.size = Pt(12)

    document.add_paragraph("")
    members_title = document.add_paragraph()
    members_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = members_title.add_run("Integrantes")
    run.bold = True
    run.font.size = Pt(12)

    for author in AUTHORS:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(author)
        run.font.size = Pt(12)

    document.add_page_break()


def add_heading(document: Document, text: str, level: int) -> None:
    document.add_heading(text, level=min(level, 2))


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def add_numbered(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Number")


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9)
        if idx < len(lines) - 1:
            run.add_break()


def add_image(document: Document, image_markdown: str) -> None:
    match = re.match(r"!\[(.*?)\]\((.*?)\)", image_markdown.strip())
    if not match:
        return

    alt_text, image_path = match.groups()
    full_path = REPORTS_DIR / image_path
    if not full_path.exists():
        return

    document.add_picture(str(full_path), width=Inches(6.2))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption = document.add_paragraph(alt_text)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].italic = True


def collect_table(lines: list[str], start_index: int) -> tuple[list[str], int]:
    table_lines = []
    i = start_index
    while i < len(lines) and "|" in lines[i]:
        table_lines.append(lines[i])
        i += 1
    return table_lines, i


def add_markdown_table(document: Document, table_lines: list[str]) -> None:
    if len(table_lines) < 2:
        return

    rows = []
    for line in table_lines:
        stripped = line.strip()
        if not stripped:
            continue
        if set(stripped.replace("|", "").replace(" ", "")) == {"-"}:
            continue
        parts = [cell.strip() for cell in stripped.strip("|").split("|")]
        rows.append(parts)

    if not rows:
        return

    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, value in enumerate(rows[0]):
        hdr_cells[idx].text = value

    for row_values in rows[1:]:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            row_cells[idx].text = value


def render_markdown(document: Document, markdown_text: str) -> None:
    lines = markdown_text.splitlines()
    i = 0
    in_code_block = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                add_code_block(document, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            add_title(document, stripped[2:].strip())
            i += 1
            continue

        if stripped.startswith("## "):
            add_heading(document, stripped[3:].strip(), 1)
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            add_numbered(document, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        if stripped.startswith("- "):
            content = stripped[2:].strip()
            image_match = re.match(r"!\[(.*?)\]\((.*?)\)", content)
            if image_match:
                add_image(document, content)
            else:
                add_bullet(document, content)
            i += 1
            continue

        if stripped.startswith("!["):
            add_image(document, stripped)
            i += 1
            continue

        if "|" in stripped:
            table_lines, new_index = collect_table(lines, i)
            add_markdown_table(document, table_lines)
            i = new_index
            continue

        document.add_paragraph(stripped)
        i += 1


def main() -> None:
    document = Document()
    set_default_font(document)
    add_cover_page(document)
    render_markdown(document, INPUT_PATH.read_text(encoding="utf-8"))
    document.save(OUTPUT_PATH)
    print(f"Reporte Word generado en: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
